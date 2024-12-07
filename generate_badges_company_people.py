import os
import csv
os.system('')

class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

location = os.getcwd() 
counter = 0

people_list = list()
# set.add(first_name + "\n" + last_name)

class Person:
    def __init__(self, first_name, last_name, company_name):
            self.first_name = first_name
            self.last_name = last_name
            self.company_name = company_name


for file in os.listdir(location):
    try:
        if file.endswith("company_people.csv"):
            print(f"{bcolors.OKGREEN}csv file found:{file}{bcolors.ENDC}")
            counter = counter + 1

            with open(file, mode='r',encoding='utf-8-sig') as csv_file:
                reader = csv.DictReader(csv_file, delimiter=',', quotechar='"')
                
                print(f"{reader.fieldnames}")


                for row in reader:
                    company_name = row['Company']
                    first_name = row['Name']
                    last_name = row['Last name']

                    person = Person(first_name, last_name, company_name)

                    people_list.append(person)

                    # if len(combined_name) > 30:
                    #     print(f"{bcolors.FAIL}LONG name:{first_name} {last_name}{bcolors.ENDC}")


    except Exception as e:
        raise e
        print("No files found here!")

print (f"{bcolors.OKCYAN}Total files found:{counter}{bcolors.ENDC}")

from docx import Document

doc = Document("badges-companies'.docx")

table = doc.tables[0]
cnt = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            try:
                if cnt == len(people_list):
                    break

                person = people_list[cnt]

                cell.text = ''
                name_run = cell.paragraphs[0].add_run(person.first_name + ' ' + person.last_name + '\n\n')
                company_run = cell.paragraphs[0].add_run(person.company_name)
                company_run.bold = True
                cnt += 1
            except Exception as e:
                raise e
                break


print (f"{bcolors.WARNING}Total names identified:{len(people_list)}{bcolors.ENDC}")
print (f"{bcolors.OKCYAN}Total names written:{cnt}{bcolors.ENDC}")


doc.save("company-people-output.docx")
