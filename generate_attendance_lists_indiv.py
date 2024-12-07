import os
import pandas as pd
from docx import Document
from docx.shared import Pt 
from docx.oxml.ns import qn  

folder_path = os.getcwd() 


def set_table_font(table, font_name="Aptos", font_size=11):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs
                for r in run:
                    r.font.name = font_name
                    r.font.size = Pt(font_size)
                    
                r_element = paragraph._element
                r_fonts = r_element.xpath(".//w:rFonts")[0]
                r_fonts.set(qn("w:eastAsia"), font_name)


for file_name in os.listdir(folder_path):
    if file_name.endswith(".csv"):
        if file_name.startswith("attendance_sheet_your_career_week_") and file_name.endswith("_speed_date.csv"):
            company_name = file_name[len("attendance_sheet_your_career_week_"):-len("_speed_date.csv")]
            company_name = company_name.replace("_", " ")

            csv_path = os.path.join(folder_path, file_name)
            df = pd.read_csv(csv_path, usecols=["First name", "Last name", "Email"])

            df['Present'] = ""

            doc = Document()
            doc.add_heading(f"Speed dates: {company_name}", level=1)
            doc.add_paragraph("Please indicate if the student showed up to the speed date with a check in the last column.")


            style = doc.styles['Normal']
            font = style.font
            font.name = "Aptos"
            font.size = Pt(12)

            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'

            for i, column_name in enumerate(df.columns):
                table.cell(0, i).text = column_name

            for row in df.itertuples(index=False):
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            set_table_font(table, font_name="Aptos", font_size=12)

            output_path = os.path.join(folder_path, f"{company_name}.docx")
            doc.save(output_path)

print("Lists generated successfully")