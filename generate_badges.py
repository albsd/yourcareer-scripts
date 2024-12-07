import os
import csv
os.system('')

class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

location = os.getcwd() 
counter = 0
otherfiles = [] 

name_set = set()
# set.add(first_name + "\n" + last_name)

for file in os.listdir(location):
    try:
        if file.endswith("students.csv"):
            print(f"{bcolors.OKGREEN}csv file found:{file}{bcolors.ENDC}")
            counter = counter + 1

            with open(file, mode='r',encoding='utf-8') as csv_file:
                reader = csv.DictReader(csv_file, delimiter=',', quotechar='"')

                for row in reader:
                    first_name = row['First name']
                    last_name = row['Last name']
                    combined_name = first_name + '\n' + last_name

                    # if len(combined_name) > 30:
                    #     print(f"{bcolors.FAIL}LONG name:{first_name} {last_name}{bcolors.ENDC}")

                    name_set.add(combined_name)



    except Exception as e:
        raise e
        print("No files found here!")

print (f"{bcolors.OKCYAN}Total files found:{counter}{bcolors.ENDC}")

from docx import Document

doc = Document("avery-docx.docx")

table = doc.tables[0]
cnt = 0
name_list = list(name_set)
name_list.sort()
iterator = iter(name_list)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            try:
                if cnt == len(name_set):
                    break

                cell.text = next(iterator)
                cnt += 1
            except Exception as e:
                raise e
                break


print (f"{bcolors.WARNING}Total names identified:{len(name_set)}{bcolors.ENDC}")
print (f"{bcolors.OKCYAN}Total names written:{cnt}{bcolors.ENDC}")


doc.save("output.docx")
